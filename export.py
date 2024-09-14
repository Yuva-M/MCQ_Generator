from docx import Document
import io

def export_to_doc(mcqs):
    doc = Document()
    for idx, mcq in enumerate(mcqs, 1):
        doc.add_heading(f"Question {idx}: {mcq['question']}", level=2)
        for choice in mcq['choices']:
            doc.add_paragraph(choice)
        doc.add_paragraph(f"Answer: {mcq['answer']}")
        doc.add_paragraph()  # Add a blank line between questions

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer